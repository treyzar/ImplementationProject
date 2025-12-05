import logging
from typing import BinaryIO
from io import BytesIO

import pdfplumber
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)


def parse_pdf_to_pages(file_obj: BinaryIO) -> list[dict]:
    """
    Парсит PDF-файл и возвращает список словарей с текстом по страницам.
    
    Args:
        file_obj: Файловый объект PDF
        
    Returns:
        Список словарей вида [{"page": 1, "text": "..."}, ...]
    """
    pages_data: list[dict] = []
    
    try:
        file_bytes = file_obj.read()
        file_obj.seek(0)
        
        with pdfplumber.open(BytesIO(file_bytes)) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):
                try:
                    text = page.extract_text() or ""
                    pages_data.append({
                        "page": page_num,
                        "text": text.strip()
                    })
                except Exception as page_error:
                    logger.warning(
                        f"Ошибка при парсинге страницы {page_num} PDF: {page_error}"
                    )
                    pages_data.append({
                        "page": page_num,
                        "text": f"[Ошибка извлечения текста со страницы {page_num}]"
                    })
                    
    except Exception as e:
        logger.error(f"Ошибка при парсинге PDF файла: {e}")
        raise ValueError(f"Не удалось распарсить PDF файл: {str(e)}")
    
    return pages_data


def parse_docx_to_pages(file_obj: BinaryIO) -> list[dict]:
    """
    Парсит DOCX-файл и возвращает список словарей с текстом.
    
    DOCX не имеет явного разделения на страницы, поэтому весь контент
    возвращается как одна "страница". При необходимости можно разбить
    по разрывам страниц или по определённому количеству параграфов.
    
    Args:
        file_obj: Файловый объект DOCX
        
    Returns:
        Список словарей вида [{"page": 1, "text": "..."}, ...]
    """
    pages_data: list[dict] = []
    
    try:
        file_bytes = file_obj.read()
        file_obj.seek(0)
        
        doc = DocxDocument(BytesIO(file_bytes))
        
        all_paragraphs: list[str] = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                all_paragraphs.append(text)
        
        for table in doc.tables:
            table_text_parts: list[str] = []
            for row in table.rows:
                row_cells: list[str] = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_cells.append(cell_text)
                if row_cells:
                    table_text_parts.append(" | ".join(row_cells))
            if table_text_parts:
                all_paragraphs.append("\n".join(table_text_parts))
        
        full_text = "\n\n".join(all_paragraphs)
        
        lines = full_text.split('\n')
        lines_per_page = 50
        
        current_page = 1
        for i in range(0, len(lines), lines_per_page):
            page_lines = lines[i:i + lines_per_page]
            page_text = "\n".join(page_lines).strip()
            if page_text:
                pages_data.append({
                    "page": current_page,
                    "text": page_text
                })
                current_page += 1
        
        if not pages_data:
            pages_data.append({
                "page": 1,
                "text": full_text.strip() if full_text.strip() else "[Документ пуст]"
            })
            
    except Exception as e:
        logger.error(f"Ошибка при парсинге DOCX файла: {e}")
        raise ValueError(f"Не удалось распарсить DOCX файл: {str(e)}")
    
    return pages_data


def get_mime_type_from_extension(filename: str) -> str:
    """
    Определяет MIME-тип по расширению файла.
    
    Args:
        filename: Имя файла
        
    Returns:
        MIME-тип
    """
    extension = filename.lower().rsplit('.', 1)[-1] if '.' in filename else ''
    
    mime_types = {
        'pdf': 'application/pdf',
        'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'doc': 'application/msword',
    }
    
    return mime_types.get(extension, 'application/octet-stream')


def parse_file_by_extension(file_obj: BinaryIO, filename: str) -> list[dict]:
    """
    Парсит файл в зависимости от его расширения.
    
    Args:
        file_obj: Файловый объект
        filename: Имя файла для определения расширения
        
    Returns:
        Список словарей с текстом по страницам
        
    Raises:
        ValueError: Если расширение не поддерживается
    """
    extension = filename.lower().rsplit('.', 1)[-1] if '.' in filename else ''
    
    if extension == 'pdf':
        return parse_pdf_to_pages(file_obj)
    elif extension == 'docx':
        return parse_docx_to_pages(file_obj)
    else:
        raise ValueError(f"Неподдерживаемое расширение файла: {extension}")
